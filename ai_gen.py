from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from groq import Groq
import streamlit as st


# ---------- LLaMA CONTENT GENERATION ----------
def generate_ai_resume(data):
    try:
        client = Groq(api_key=st.secrets["GROQ_API_KEY"])

        prompt = f"""
        Create a professional ATS-friendly resume content.

        Name: {data.get("name")}
        Skills: {data.get("skills")}
        Education: {data.get("education")}
        Experience: {data.get("experience")}

        Format:
        - Skills as bullet points
        - Experience as strong action statements
        - Keep it concise and professional
        """

        response = client.chat.completions.create(
            model="llama3-70b-8192",
            messages=[{"role": "user", "content": prompt}],
        )

        return response.choices[0].message.content

    except Exception as e:
        return None  # fallback


# ---------- DOCX ----------
def create_docx(data):
    doc = Document()

    ai_text = generate_ai_resume(data)

    # Name
    doc.add_heading(data.get("name", ""), 0)

    contact = f"{data.get('email','')} | {data.get('phone','')}"
    if data.get("links"):
        contact += f" | {data.get('links')}"
    doc.add_paragraph(contact)

    # If AI worked → use it
    if ai_text:
        for line in ai_text.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
    else:
        # fallback manual formatting
        doc.add_heading("Skills", level=1)
        for skill in data.get("skills", "").split(","):
            if skill.strip():
                doc.add_paragraph(skill.strip(), style="List Bullet")

        doc.add_heading("Education", level=1)
        doc.add_paragraph(data.get("education", ""))

        doc.add_heading("Experience", level=1)
        for line in data.get("experience", "").split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip(), style="List Bullet")

    file_path = "resume.docx"
    doc.save(file_path)
    return file_path


# ---------- PDF ----------
def create_pdf(data):
    doc = SimpleDocTemplate("resume.pdf")
    styles = getSampleStyleSheet()
    elements = []

    ai_text = generate_ai_resume(data)

    # Name
    elements.append(Paragraph(f"<b>{data.get('name','')}</b>", styles['Title']))
    elements.append(Spacer(1, 10))

    contact = f"{data.get('email','')} | {data.get('phone','')}"
    if data.get("links"):
        contact += f" | {data.get('links')}"
    elements.append(Paragraph(contact, styles['Normal']))
    elements.append(Spacer(1, 12))

    if ai_text:
        for line in ai_text.split("\n"):
            if line.strip():
                elements.append(Paragraph(line.strip(), styles['Normal']))
    else:
        elements.append(Paragraph("<b>Skills</b>", styles['Heading2']))
        for skill in data.get("skills", "").split(","):
            if skill.strip():
                elements.append(Paragraph(f"• {skill.strip()}", styles['Normal']))

        elements.append(Spacer(1, 12))

        elements.append(Paragraph("<b>Education</b>", styles['Heading2']))
        elements.append(Paragraph(data.get("education",""), styles['Normal']))

        elements.append(Spacer(1, 12))

        elements.append(Paragraph("<b>Experience</b>", styles['Heading2']))
        for line in data.get("experience","").split("\n"):
            if line.strip():
                elements.append(Paragraph(f"• {line.strip()}", styles['Normal']))

    doc.build(elements)
    return "resume.pdf"
