from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet


# ---------- DOCX ----------
def create_docx(data):
    doc = Document()

    doc.add_heading(data["name"], 0)

    doc.add_paragraph(f"{data['email']} | {data['phone']}")
    doc.add_paragraph(f"{data['linkedin']} | {data['github']}")

    doc.add_heading("Summary", 1)
    doc.add_paragraph(data["summary"])

    doc.add_heading("Skills", 1)
    for s in data["skills"].split(","):
        if s.strip():
            doc.add_paragraph(s.strip(), style="List Bullet")

    doc.add_heading("Experience", 1)
    for e in data["experience"].split("\n"):
        if e.strip():
            doc.add_paragraph(e.strip(), style="List Bullet")

    doc.add_heading("Projects", 1)
    for p in data["projects"].split("\n"):
        if p.strip():
            doc.add_paragraph(p.strip(), style="List Bullet")

    doc.add_heading("Education", 1)
    doc.add_paragraph(data["education"])

    doc.add_heading("Achievements", 1)
    doc.add_paragraph(data["achievements"])

    doc.save("resume.docx")
    return "resume.docx"


# ---------- PDF ----------
def create_pdf(data):
    doc = SimpleDocTemplate("resume.pdf")
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph(f"<b>{data['name']}</b>", styles["Title"]))
    elements.append(Spacer(1, 10))

    elements.append(Paragraph(data["summary"], styles["Normal"]))
    elements.append(Spacer(1, 12))

    for section in ["skills", "experience", "projects"]:
        elements.append(Paragraph(f"<b>{section.title()}</b>", styles["Heading2"]))
        for item in data[section].split("\n"):
            if item.strip():
                elements.append(Paragraph(f"• {item}", styles["Normal"]))
        elements.append(Spacer(1, 10))

    doc.build(elements)
    return "resume.pdf"


# ---------- PORTFOLIO ----------
def create_portfolio(data):
    html = f"""
<html>
<head>
<style>
body {{
    font-family: 'Segoe UI', sans-serif;
    margin: 0;
    background: #0f172a;
    color: #e2e8f0;
}}

.hero {{
    text-align: center;
    padding: 80px;
    background: linear-gradient(135deg, #1e293b, #334155);
}}

.section {{
    padding: 50px 15%;
}}

.card {{
    background: #1e293b;
    padding: 20px;
    margin: 15px 0;
    border-radius: 10px;
}}

h2 {{
    color: #38bdf8;
}}
</style>
</head>

<body>

<div class="hero">
    <h1>{data['name']}</h1>
    <h3>{data['title']}</h3>
    <p>{data['email']} | {data['phone']}</p>
</div>

<div class="section">
    <h2>About</h2>
    <p>{data['summary']}</p>
</div>

<div class="section">
    <h2>Projects</h2>
    {"".join([f"<div class='card'>{p}</div>" for p in data['projects'].split("\\n") if p.strip()])}
</div>

<div class="section">
    <h2>Experience</h2>
    {"".join([f"<div class='card'>{e}</div>" for e in data['experience'].split("\\n") if e.strip()])}
</div>

</body>
</html>
"""

    with open("portfolio.html", "w", encoding="utf-8") as f:
        f.write(html)

    return "portfolio.html"
